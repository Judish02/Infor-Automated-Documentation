import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

src = '01_Reference_Examples/002_IN_A/DES-030_002_IN_A_REFLEX_Préparation_Cde_V1.0.docx'
out = '02_Current_Target/DES030_5058_IE-TDX-01-CDV_Orders_V1_0.docx'

doc = Document(src)

# ── CONFIRMED COMPONENT FILE NAMES ──────────────────────────────────────────
ZAP_MEC    = 'M3EDI_M3EDISalesOrderCustom_Load_In_2_13_0.zap'
ZIP_SCHEMA = 'M3EDISalesOrderCustom.zip'
XML_MAP    = 'M3EDISalesOrderCustomLoadMapping.xml'
ZIP_NOTIF  = 'M3IECNotification.zip'
XML_DF     = 'Spoon_LoadM3EDISalesOrderCustom.xml'
XML_MON    = 'Spoon_M3EDISalesOrderCustomNotification.xml'

BOD_LOAD   = 'Load.M3EDISalesOrderCustom'
BOD_ACK    = 'Acknowledge.M3EDISalesOrderCustom'
BOD_LOAD_S = 'LoadM3EDISalesOrderCustom'
BOD_ACK_S  = 'AcknowledgeM3EDISalesOrderCustom'
AGREEMENT  = 'M3_M3EDISalesOrderCustom_Load_In'
SCHEMA_NM  = 'M3EDISalesOrderCustom'
SCHEMA_NM2 = 'M3IECNotification'

DIR_IN     = '/PHEERPAPP01R (Interfaces)/EDI/VENTE/CF0/ORDERS'
DIR_ERR    = '/PHEERPAPP01R (Interfaces)/EDI/VENTE/CF0/ORDERS/ERREUR'
DIR_ARC    = '/PHEERPAPP01R (Interfaces)/EDI/VENTE/CF0/ORDERS/ARCH'
FILE_PAT   = 'ORDERS-IN_*.xml'

# ── HELPERS ─────────────────────────────────────────────────────────────────
def set_para(para, text):
    for r in para.runs:
        r.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

def replace_in_para(para, old, new):
    if old not in para.text:
        return False
    rebuilt = para.text.replace(old, new)
    set_para(para, rebuilt)
    return True

def replace_all(old, new):
    for p in doc.paragraphs:
        replace_in_para(p, old, new)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_para(p, old, new)
    for sec in doc.sections:
        for p in sec.footer.paragraphs + sec.header.paragraphs:
            replace_in_para(p, old, new)

def set_cell(cell, text):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
    else:
        cell.paragraphs[0].add_run(text)

# ── STEP 1: GLOBAL REPLACEMENTS (most specific first) ───────────────────────
ordered = [
    # Old MEC Mapper .zap pattern from template
    ('CSTM3BOD_REFLEXPreparationCdeCustom_Acknowledge_Error_Out_1_0_0', ZAP_MEC.replace('.zap','')),
    ('CSTM3BOD_REFLEXPreparationCdeCustom_Process_In_1_0_0',            ZAP_MEC.replace('.zap','')),
    # Agreement name
    ('M3_REFLEXPreparationCdeCustom_Process_In', AGREEMENT),
    # BOD dotted names
    ('Process.REFLEXPreparationCdeCustom',      BOD_LOAD),
    ('Acknowledge.REFLEXPreparationCdeCustom',  BOD_ACK),
    # File pattern
    ('ProcessREFLEXPreparationCdeCustom_*.xml', FILE_PAT),
    # Concatenated BOD names
    ('ProcessREFLEXPreparationCdeCustom',       BOD_LOAD_S),
    ('AcknowledgeREFLEXPreparationCdeCustom',   BOD_ACK_S),
    # Schema base name
    ('REFLEXPreparationCdeCustom',              SCHEMA_NM),
    # Interface identity
    ('002_IN_A REFLEX Préparation Cde',         '5058 IE-TDX-01-CDV Orders'),
    ('002_IN_A REFLEX Préparation_Cde',         '5058 IE-TDX-01-CDV Orders'),
    ('002_IN_A REFLEX',                         '5058 IE-TDX-01'),
    ('002_IN_A',                                '5058'),
    ('REFLEX -> M3',                            'TDX Infinity -> M3'),
    ('REFLEX',                                  'TDX'),
    # Author & dates
    ('Harshil Tonoo', 'Motean Kenylen'),
    ('Harshil',       'Motean Kenylen'),
    ('27/03/2025',    '25/05/26'),
    # Object Schema zip reference
    (SCHEMA_NM + '.zip', ZIP_SCHEMA),
    # Script file
    ('SC_002_IN_A.json', 'SC_5058_IN_A.json'),
    ('SC_002_IN_A',      'SC_5058_IN_A'),
    # MBDREADMI config
    ('ConfigData_MDBREADMI_GETOCUSMA00_20250327.xml', 'ConfigData_MDBREADMI_GETOCUSMA00_5058.xml'),
]
for old, new in ordered:
    replace_all(old, new)
print("Global replacements done.")

# ── STEP 2: TABLE 0 — Document Properties ───────────────────────────────────
t0 = doc.tables[0]
set_cell(t0.rows[1].cells[1], 'Motean Kenylen')
set_cell(t0.rows[2].cells[1], '25/05/26')
set_cell(t0.rows[3].cells[1], '25/05/26')
set_cell(t0.rows[4].cells[1], '5058-IE-TDX-01-CDV-ORDERS')
set_cell(t0.rows[5].cells[1], '1.0')
print("Table 0 done.")

# ── STEP 3: TABLE 2 — Change Record ─────────────────────────────────────────
t2 = doc.tables[2]
set_cell(t2.rows[1].cells[0], '25/05/26')
set_cell(t2.rows[1].cells[1], 'Motean Kenylen')
set_cell(t2.rows[1].cells[2], '1.0')
set_cell(t2.rows[1].cells[3], 'Creation du document')
for ri in range(2, len(t2.rows)):
    for ci in range(4):
        set_cell(t2.rows[ri].cells[ci], '')
print("Table 2 done.")

# ── STEP 4: TABLE 4 — Deliverables ──────────────────────────────────────────
t4 = doc.tables[4]
rows_data = [
    ['Custom BOD / Agreement', 'M3EDI_M3EDISalesOrderCustom_Load_In',
     'M3EDI_M3EDISalesOrderCustom_Load_In', ZAP_MEC],
    ['Object Schema', SCHEMA_NM, SCHEMA_NM, ZIP_SCHEMA],
    ['ION Mapping', 'M3EDISalesOrderCustomLoadMapping',
     'M3EDISalesOrderCustomLoadMapping', XML_MAP],
    ['Object Schema (Notification)', SCHEMA_NM2, SCHEMA_NM2, ZIP_NOTIF],
    ['Document Flow', 'Spoon_LoadM3EDISalesOrderCustom',
     'Spoon_LoadM3EDISalesOrderCustom', XML_DF],
    ['Monitor', 'Spoon_M3EDISalesOrderCustomNotification',
     'Spoon_M3EDISalesOrderCustomNotification', XML_MON],
]
for ri, rd in enumerate(rows_data):
    idx = ri + 1
    if idx < len(t4.rows):
        for ci, val in enumerate(rd):
            if ci < len(t4.rows[idx].cells):
                set_cell(t4.rows[idx].cells[ci], val)
for ri in range(len(rows_data) + 1, len(t4.rows)):
    for cell in t4.rows[ri].cells:
        set_cell(cell, '')
print("Table 4 done.")

# ── STEP 5: TABLE OF CONTENTS — update section titles ───────────────────────
toc_updates = {
    '3.4\tScript': '3.4\tION Mapping',
    '3.6\tMBDREADMI': '3.6\tObject Schema (Notification)',
    '3.8\tTranscodification': '3.8\tMonitor',
}
for p in doc.paragraphs:
    for old, new in toc_updates.items():
        if old in p.text:
            set_para(p, p.text.replace(old, new))
print("TOC entries updated.")

# ── STEP 6: INTRODUCTION ────────────────────────────────────────────────────
for p in doc.paragraphs:
    if 'informations nécessaires' in p.text and ('promotion' in p.text or 'déploiement' in p.text):
        set_para(p, ("Ce document identifie les informations nécessaires au déploiement de "
                     "l'interface 5058 IE-TDX-01-CDV Orders de l'environnement DEV vers "
                     "les environnements TEST/PROD."))
        break
print("Introduction done.")

# ── STEP 7: OBJECT SCHEMA section — update filename ─────────────────────────
for p in doc.paragraphs:
    if 'Importation du custom Objet Schema' in p.text:
        set_para(p, '3.2.1 Importation du custom Objet Schema ' + SCHEMA_NM)
    if 'Localisez le fichier' in p.text and 'zip' in p.text and SCHEMA_NM in p.text:
        set_para(p, 'Localisez le fichier ' + ZIP_SCHEMA + ' et cliquez OK')
print("Object Schema section done.")

# ── STEP 8: CONNECTION POINT — specific paths ────────────────────────────────
for p in doc.paragraphs:
    t = p.text
    if "Emplacement de lecture" in t and "entrant" in t:
        set_para(p, "Dans la case 'Emplacement de lecture', renseignez : " + DIR_IN)
    elif "Emplacement d" in t and "erreur" in t and "répertoire" in t:
        set_para(p, "Dans la case 'Emplacement d'erreur', renseignez : " + DIR_ERR)
    elif "Emplacement d'archive" in t:
        set_para(p, "Dans la case 'Emplacement d'archive', renseignez : " + DIR_ARC)
    elif "Modèle de nom de fichier" in t and "ORDERS" in t:
        set_para(p, "Dans la case 'Modèle de nom de fichier', renseignez le nom du fichier -> '" + FILE_PAT + "'")
print("Connection point paths done.")

# ── STEP 9: SCRIPT section → repurpose as ION MAPPING ───────────────────────
script_heading_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and 'Script' in p.text and 'Document' not in p.text:
        script_heading_idx = i
        set_para(p, 'ION Mapping')
        break

if script_heading_idx is not None:
    # Repurpose body paragraphs in Script section for ION Mapping steps
    ion_map_steps = [
        '',
        '3.4.1 Importation du ION Mapping M3EDISalesOrderCustomLoadMapping',
        'Naviguez dans ION > Connect > Mappages > Importer',
        '',
        '',
        'Sélectionnez le fichier à importer ' + XML_MAP,
        '',
        'Cliquez sur OK',
        '',
        "Vérifiez que le mapping 'M3EDISalesOrderCustomLoadMapping' est présent dans la liste",
        '',
        '',
    ]
    # Find the body paragraphs that follow the Script heading until next Heading 2
    para_list = doc.paragraphs
    body_idx = script_heading_idx + 1
    step_i = 0
    for j in range(body_idx, min(body_idx + 15, len(para_list))):
        p = para_list[j]
        if p.style.name.startswith('Heading') and j > body_idx:
            break
        if step_i < len(ion_map_steps):
            set_para(p, ion_map_steps[step_i])
            step_i += 1
    print("ION Mapping section set.")

# ── STEP 10: DOCUMENT FLOW — update DF filename ─────────────────────────────
for p in doc.paragraphs:
    t = p.text
    if 'Sélectionnez le fichier à importer' in t and ('DF_' in t or 'Spoon_' in t or '5058' in t):
        set_para(p, 'Sélectionnez le fichier à importer ' + XML_DF)
    elif 'DF_5058' in t or 'DF_002' in t:
        set_para(p, p.text.replace('DF_5058_IN_A_TDX.xml', XML_DF)
                         .replace('DF_5058', XML_DF.replace('.xml', ''))
                         .replace('DF_002', XML_DF.replace('.xml', '')))
print("Document Flow filename done.")

# ── STEP 11: MEC MAPPER — update .zap filename ──────────────────────────────
# The global replacements already set both heading 3 zap names to ZAP_MEC base.
# Now fix: there should be ONE zap (not two). Adjust heading 3 texts.
zap_base = ZAP_MEC.replace('.zap', '')
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if p.style.name == 'Heading 3' and zap_base in t:
        # Deduplicate: first occurrence = Load_In publication
        # We keep only one MEC Mapper section (Load_In)
        if "Publication" in t:
            set_para(p, "Publication de l'agreement '" + zap_base + "'")
            break

# Fix Browse step referencing the .zap file
for p in doc.paragraphs:
    if "Browse" in p.text and zap_base in p.text and '.zap' not in p.text:
        set_para(p, p.text.replace(zap_base, ZAP_MEC))
    elif "Browse" in p.text and 'M3EDISalesOrderCustom' in p.text:
        set_para(p, "Cliquez sur « Browse » et sélectionnez le fichier " + ZAP_MEC + " et cliquez « OK ».")

# For the "Etendez le" paragraphs
for p in doc.paragraphs:
    if 'Etendez' in p.text and zap_base in p.text:
        set_para(p, "Etendez le « " + zap_base + " » et faites clique droite sur Documents.")

# Validate / Publish steps
for p in doc.paragraphs:
    if ('Faites clique' in p.text or 'Faites clique' in p.text) and ('Validate' in p.text or 'Publish' in p.text) and zap_base in p.text:
        pass  # already correctly replaced

# Agreement activation
for p in doc.paragraphs:
    if 'agreement' in p.text.lower() and 'Activate' in p.text and AGREEMENT in p.text:
        set_para(p, ("Cherchez l'agreement « " + AGREEMENT + " ». "
                     "Cliquez sur le bouton « Activate » pour activer le agreement."))
print("MEC Mapper section done.")

# ── STEP 12: MBDREADMI section → repurpose as Object Schema (Notification) ──
mbdr_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and 'MBDREADMI' in p.text:
        mbdr_idx = i
        set_para(p, 'Object Schema (Notification)')
        break

if mbdr_idx is not None:
    notif_steps = [
        '',
        '3.6.1 Importation du custom Objet Schema ' + SCHEMA_NM2,
        'Allez dans OS > ION > Catalogue de données > Schémas d objet',
        '',
        '',
        'Cliquez sur le bouton Importer',
        '',
        '',
        'Localisez le fichier ' + ZIP_NOTIF + ' et cliquez OK',
        '',
        '',
        "Vérifiez que le schéma '" + SCHEMA_NM2 + "' est présent dans la liste.",
        '',
        '',
        "Cliquez sur 'Activer' pour activer le schéma.",
        '',
        '',
        '',
        '',
        '',
    ]
    para_list = doc.paragraphs
    body_idx = mbdr_idx + 1
    step_i = 0
    for j in range(body_idx, min(body_idx + 50, len(para_list))):
        p = para_list[j]
        if p.style.name.startswith('Heading') and j > body_idx:
            break
        if step_i < len(notif_steps):
            set_para(p, notif_steps[step_i])
            step_i += 1
        else:
            set_para(p, '')
    print("Object Schema Notification section set.")

# ── STEP 13: TRANSCODIFICATION section → repurpose as MONITOR ───────────────
trans_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and 'Transcodification' in p.text:
        trans_idx = i
        set_para(p, 'Monitor')
        break

if trans_idx is not None:
    # Fix Heading 3 below
    para_list = doc.paragraphs
    for j in range(trans_idx + 1, min(trans_idx + 5, len(para_list))):
        p = para_list[j]
        if p.style.name == 'Heading 3':
            set_para(p, "Configuration du Monitor Spoon_M3EDISalesOrderCustomNotification")
            break

    monitor_steps = [
        '',
        '3.8.1 Importation et configuration du Monitor de notification',
        'Naviguez dans ION > Connect > Moniteurs',
        '',
        '',
        'Cliquez sur le bouton Importer',
        '',
        '',
        'Sélectionnez le fichier ' + XML_MON + ' et cliquez OK',
        '',
        '',
        "Vérifiez que le monitor 'Spoon_M3EDISalesOrderCustomNotification' est présent dans la liste.",
        '',
        "Activez le monitor en cliquant sur le bouton 'Activer'.",
        '',
        "Configurez les destinataires de notification selon les paramètres CRS424 définis dans le DES-020.",
    ]
    body_idx = trans_idx + 1
    step_i = 0
    # Skip the heading 3
    for j in range(body_idx, min(body_idx + 20, len(para_list))):
        p = para_list[j]
        if p.style.name == 'Heading 3':
            continue
        if step_i < len(monitor_steps):
            set_para(p, monitor_steps[step_i])
            step_i += 1
        else:
            set_para(p, '')
    print("Monitor section set.")

# ── STEP 14: TITLE BLOCK ────────────────────────────────────────────────────
p1 = doc.paragraphs[1]
set_para(p1, '5058 IE-TDX-01-CDV Orders')
p2 = doc.paragraphs[2]
set_para(p2, 'TDX Infinity -> M3')
print("Title block done.")

# ── STEP 15: FOOTER ─────────────────────────────────────────────────────────
for sec in doc.sections:
    for p in sec.footer.paragraphs:
        for old, new in [
            ('002_IN_A REFLEX', '5058 IE-TDX-01-CDV'),
            ('TDX Préparation Cde', '5058 IE-TDX-01-CDV Orders'),
            ('Préparation Cde', 'IE-TDX-01-CDV Orders'),
            ('DES-030_5058_TDX_Préparation_Cde_V1.0', 'DES-030_5058_IE-TDX-01-CDV_Orders_V1.0'),
            ('002', '5058'),
            ('REFLEX', 'TDX'),
        ]:
            replace_in_para(p, old, new)
print("Footer done.")

# ── SAVE ────────────────────────────────────────────────────────────────────
doc.save(out)
print("\nSaved:", out)

# ── VERIFY ──────────────────────────────────────────────────────────────────
doc2 = Document(out)
print("\n=== VERIFICATION ===")
print("Title P1:", doc2.paragraphs[1].text)
print("Title P2:", doc2.paragraphs[2].text)
print("\nTable 0:")
for row in doc2.tables[0].rows:
    print(" ", row.cells[0].text[:30], ":", row.cells[1].text[:50])
print("\nTable 4 (Deliverables):")
for row in doc2.tables[4].rows:
    print("  Type:", row.cells[0].text[:40], "| File:", row.cells[3].text[:60])
print("\nKey section headings:")
for p in doc2.paragraphs:
    if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3') and p.text.strip():
        print(" ", p.style.name, ":", p.text[:80])
print("\nFooter:")
for p in doc2.sections[0].footer.paragraphs:
    if p.text.strip():
        print(" ", p.text[:80])
