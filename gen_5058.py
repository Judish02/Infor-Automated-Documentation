import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

src  = '01_Reference_Examples/002_IN_A/DES-030_002_IN_A_REFLEX_Préparation_Cde_V1.0.docx'
out  = '02_Current_Target/DES030_5058_IE-TDX-01-CDV_Orders_V1_0.docx'

doc = Document(src)

# ──────────────────────────────────────────────
# HELPER
# ──────────────────────────────────────────────
def replace_in_para(para, old, new):
    full = para.text
    if old not in full:
        return False
    rebuilt = full.replace(old, new)
    for r in para.runs:
        r.text = ''
    if para.runs:
        para.runs[0].text = rebuilt
    return True

def replace_all(old, new):
    for p in doc.paragraphs:
        replace_in_para(p, old, new)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_para(p, old, new)
    for sec in doc.sections:
        for p in sec.footer.paragraphs:
            replace_in_para(p, old, new)
        for p in sec.header.paragraphs:
            replace_in_para(p, old, new)

def set_cell(cell, text):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
    else:
        cell.paragraphs[0].add_run(text)

# ──────────────────────────────────────────────
# ORDERED REPLACEMENTS
# ──────────────────────────────────────────────
replacements = [
    ('CSTM3BOD_REFLEXPreparationCdeCustom_Acknowledge_Error_Out_1_0_0',
     'CSTM3BOD_M3EDISalesOrderCustom_Acknowledge_Error_Out_1_0_0'),
    ('CSTM3BOD_REFLEXPreparationCdeCustom_Process_In_1_0_0',
     'CSTM3BOD_M3EDISalesOrderCustom_Load_In_1_0_0'),
    ('M3_REFLEXPreparationCdeCustom_Process_In',
     'M3_M3EDISalesOrderCustom_Load_In'),
    ('Process.REFLEXPreparationCdeCustom',
     'Load.M3EDISalesOrderCustom'),
    ('Acknowledge.REFLEXPreparationCdeCustom',
     'Acknowledge.M3EDISalesOrderCustom'),
    ('ProcessREFLEXPreparationCdeCustom_*.xml',
     'ORDERS-IN_*.xml'),
    ('ProcessREFLEXPreparationCdeCustom',
     'LoadM3EDISalesOrderCustom'),
    ('AcknowledgeREFLEXPreparationCdeCustom',
     'AcknowledgeM3EDISalesOrderCustom'),
    ('REFLEXPreparationCdeCustom',
     'M3EDISalesOrderCustom'),
    ('002_IN_A REFLEX Préparation Cde',
     '5058 IE-TDX-01-CDV Orders'),
    ('002_IN_A REFLEX Préparation_Cde',
     '5058 IE-TDX-01-CDV Orders'),
    ('002_IN_A REFLEX',
     '5058 IE-TDX-01'),
    ('002_IN_A',
     '5058'),
    ('REFLEX -> M3',
     'TDX Infinity -> M3'),
    ('REFLEX',
     'TDX'),
    ('Harshil Tonoo', 'Motean Kenylen'),
    ('Harshil',       'Motean Kenylen'),
    ('27/03/2025', '25/05/26'),
    ('ConfigData_MDBREADMI_GETOCUSMA00_20250327.xml',
     'ConfigData_MDBREADMI_GETOCUSMA00_5058.xml [A CONFIRMER]'),
    ('SC_002_IN_A.json', 'SC_5058_IN_A.json [A CONFIRMER]'),
    ('SC_002_IN_A',      'SC_5058_IN_A'),
]

for old, new in replacements:
    replace_all(old, new)
print("Global replacements done.")

# ──────────────────────────────────────────────
# TABLE 0
# ──────────────────────────────────────────────
t0 = doc.tables[0]
set_cell(t0.rows[1].cells[1], 'Motean Kenylen')
set_cell(t0.rows[2].cells[1], '25/05/26')
set_cell(t0.rows[3].cells[1], '25/05/26')
set_cell(t0.rows[4].cells[1], '5058-IE-TDX-01-CDV-ORDERS')
set_cell(t0.rows[5].cells[1], '1.0')
print("Table 0 updated.")

# ──────────────────────────────────────────────
# TABLE 2
# ──────────────────────────────────────────────
t2 = doc.tables[2]
set_cell(t2.rows[1].cells[0], '25/05/26')
set_cell(t2.rows[1].cells[1], 'Motean Kenylen')
set_cell(t2.rows[1].cells[2], '1.0')
set_cell(t2.rows[1].cells[3], 'Creation du document')
for ri in range(2, len(t2.rows)):
    for ci in range(4):
        set_cell(t2.rows[ri].cells[ci], '')
print("Table 2 updated.")

# ──────────────────────────────────────────────
# TABLE 4: Deliverables
# ──────────────────────────────────────────────
t4 = doc.tables[4]
new_rows = [
    ['Object Schema',
     'M3EDISalesOrderCustom',
     'M3EDISalesOrderCustom',
     'M3EDISalesOrderCustom.zip'],
    ['Document Flow',
     'DF_5058_IN_A_TDX',
     'DF_5058_IN_A_TDX',
     'DF_5058_IN_A_TDX.xml [A CONFIRMER]'],
    ['MBDREADMI - Transactions',
     'GetOCUSMA00 [A CONFIRMER]',
     'GetOCUSMA00',
     'ConfigData_MDBREADMI_GETOCUSMA00_5058.xml [A CONFIRMER]'],
    ['Mec Mapper - Acknowledge BOD',
     'CSTM3BOD_M3EDISalesOrderCustom_Acknowledge_Error_Out_1_0_0',
     'CSTM3BOD_M3EDISalesOrderCustom_Acknowledge_Error_Out_1_0_0',
     'CSTM3BOD_M3EDISalesOrderCustom_Acknowledge_Error_Out_1_0_0.zap [A CONFIRMER]'],
    ['Mec Mapper',
     'CSTM3BOD_M3EDISalesOrderCustom_Load_In_1_0_0',
     'CSTM3BOD_M3EDISalesOrderCustom_Load_In_1_0_0',
     'CSTM3BOD_M3EDISalesOrderCustom_Load_In_1_0_0.zap [A CONFIRMER]'],
]
for ri, row_data in enumerate(new_rows):
    idx = ri + 1
    if idx < len(t4.rows):
        row = t4.rows[idx]
        for ci, val in enumerate(row_data):
            if ci < len(row.cells):
                set_cell(row.cells[ci], val)
for ri in range(len(new_rows) + 1, len(t4.rows)):
    for cell in t4.rows[ri].cells:
        set_cell(cell, '')
print("Table 4 updated.")

# ──────────────────────────────────────────────
# INTRODUCTION
# ──────────────────────────────────────────────
for p in doc.paragraphs:
    t = p.text
    if 'informations nécessaires' in t and ('promotion' in t or 'déploiement' in t):
        rebuilt = ("Ce document identifie les informations nécessaires au déploiement de "
                   "l'interface 5058 IE-TDX-01-CDV Orders de l'environnement DEV vers "
                   "les environnements TEST/PROD.")
        for r in p.runs:
            r.text = ''
        if p.runs:
            p.runs[0].text = rebuilt
        print("Introduction updated.")
        break

# ──────────────────────────────────────────────
# CONNECTION POINT: specific directory paths
# ──────────────────────────────────────────────
for p in doc.paragraphs:
    t = p.text
    if "Emplacement de lecture" in t and "entrant" in t:
        rebuilt = ("Dans la case 'Emplacement de lecture', renseignez : "
                   "/PHEERPAPP01R (Interfaces)/EDI/VENTE/CF0/ORDERS")
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = rebuilt
    elif "Emplacement d" in t and "erreur" in t and "répertoire" in t:
        rebuilt = ("Dans la case 'Emplacement d'erreur', renseignez : "
                   "/PHEERPAPP01R (Interfaces)/EDI/VENTE/CF0/ORDERS/ERREUR")
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = rebuilt
    elif "Emplacement d'archive" in t:
        rebuilt = ("Dans la case 'Emplacement d'archive', renseignez : "
                   "/PHEERPAPP01R (Interfaces)/EDI/VENTE/CF0/ORDERS/ARCH")
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = rebuilt
print("Connection point paths updated.")

# ──────────────────────────────────────────────
# SCRIPT SECTION: mark as N/A
# ──────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and p.text.strip() == 'Script [A CONFIRMER]':
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = 'Script [Non applicable pour cette interface]'
        break
    elif p.style.name == 'Heading 2' and 'Script' in p.text and 'Document' not in p.text and 'flow' not in p.text.lower():
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = 'Script [Non applicable pour cette interface]'
        break

for p in doc.paragraphs:
    t = p.text
    if 'SC_5058_IN_A' in t and ('json' in t.lower() or 'CONFIRMER' in t):
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = "N/A - Cette interface n'utilise pas de composant Script."

# DataFlow Script activity reference
for p in doc.paragraphs:
    t = p.text
    if 'activité Script' in t and 'Propriété' in t:
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = ("[N/A] Cette interface ne contient pas d'activité Script dans le "
                              "DataFlow. Vérifiez les activités selon la configuration DES-020 5058.")
print("Script section marked N/A.")

# ──────────────────────────────────────────────
# TITLE BLOCK P1 & P2
# ──────────────────────────────────────────────
paras = doc.paragraphs
# Row 1 = interface label, Row 2 = direction
p1 = paras[1]
for r in p1.runs: r.text = ''
if p1.runs: p1.runs[0].text = '5058 IE-TDX-01-CDV Orders'
p2 = paras[2]
for r in p2.runs: r.text = ''
if p2.runs: p2.runs[0].text = 'TDX Infinity -> M3'
print("Title block updated.")

# ──────────────────────────────────────────────
# FOOTER
# ──────────────────────────────────────────────
for sec in doc.sections:
    for p in sec.footer.paragraphs:
        replace_in_para(p, '002_IN_A REFLEX', '5058 IE-TDX-01-CDV')
        replace_in_para(p, '002_IN_A', '5058')
        replace_in_para(p, 'REFLEX', 'TDX')
print("Footer updated.")

doc.save(out)
print("\nSaved:", out)

# Quick verify
doc2 = Document(out)
print("P1:", doc2.paragraphs[1].text)
print("P2:", doc2.paragraphs[2].text)
print("T0 Author:", doc2.tables[0].rows[1].cells[1].text)
print("T0 Date:", doc2.tables[0].rows[2].cells[1].text)
print("T0 Ref:", doc2.tables[0].rows[4].cells[1].text)
print("T4 rows:", len(doc2.tables[4].rows))
for r in doc2.tables[4].rows:
    print("  ", [c.text[:40] for c in r.cells])
